#!/usr/bin/env python3
"""
build_brief.py - the VIDEN creative-brief / script-doc .docx builder.

The plugin's single document builder. Takes a structured brief (JSON or a small
markdown dialect) and emits a .docx styled to the VIDEN client-document system
described in ../references/docx-style-spec.md.

Usage
-----
    python3 build_brief.py --in brief.json --out brief.docx
    python3 build_brief.py --in brief.md   --out brief.docx
    python3 build_brief.py --self-test --out /tmp/selftest.docx
    python3 build_brief.py --in brief.json --out brief.docx --fonts ./Urbanist --embed

Input: JSON
-----------
    {
      "title": "Client name - creative brief",
      "subtitle": "August 2026",
      "blocks": [
        {"type": "callout", "label": "Executive summary", "body": ["one", "two"]},
        {"type": "h2", "text": "Concept 1 - the reframe"},
        {"type": "sublabel", "text": "Strategy"},
        {"type": "fields", "items": [["Angle", "..."], ["Persona", "..."]]},
        {"type": "body", "text": "A paragraph."},
        {"type": "bullets", "items": [{"lead": "Lead", "text": " continuation"},
                                      "a plain bullet"]},
        {"type": "table", "columns": ["Line", "Visual direction"],
         "widths": [4680, 4680], "align": ["left", "left"],
         "rows": [["...", "..."]]},
        {"type": "spacer"},
        {"type": "pagebreak"}
      ]
    }

Input: markdown dialect
-----------------------
    # Title
    Single line straight after the title becomes the subtitle.
    ## Section header            -> h2
    ### Sub label                -> sublabel
    > **Executive summary**      -> callout label; following > lines are its body
    **Angle:** value             -> field row (consecutive lines group together)
    - **Lead** continuation      -> bullet with a bold navy lead-in
    | a | b |                    -> table (the |---| separator row is optional)
    ---                          -> spacer
    anything else                -> body paragraph

Blank lines separate blocks. Em dashes, en dashes and minus signs are converted
to short hyphens on output, per the house rule.

Fonts
-----
Urbanist is native in Google Docs and is the document font. Embedding is only
needed when the .docx must render correctly on a machine without Urbanist
installed. Pass --fonts DIR --embed with a directory containing
Urbanist-Regular.ttf / -SemiBold.ttf / -Italic.ttf / -SemiBoldItalic.ttf
(any subset; whatever is present gets embedded). Without --embed the file is
still valid and correct, it just falls back to Montserrat then Arial off-brand.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
import uuid
import zipfile
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# --------------------------------------------------------------------------
# Style constants - the single source of these values is
# ../references/docx-style-spec.md. Change them there first.
# --------------------------------------------------------------------------

FONT = "Urbanist"
FONT_FALLBACKS = ("Montserrat", "Arial")

NAVY = "202938"        # titles, section headers, table-header fill, big numbers
BRAND_BLUE = "3B63F6"  # section marker, arrows, hyperlinks, callout label
GRAPHITE = "2C313A"    # body text
MUTED = "6B7280"       # labels, captions, source notes
POSITIVE = "1B8A5A"    # improved
NEGATIVE = "C0392B"    # worsened
CALLOUT_FILL = "EEF2FB"
CALLOUT_BORDER = "D7E0F5"
ROW_HAIRLINE = "ECECEC"
WHITE = "FFFFFF"

PAGE_W_DXA = 12240     # US Letter
PAGE_H_DXA = 15840
MARGIN_DXA = 1440      # 1 inch
CONTENT_DXA = PAGE_W_DXA - 2 * MARGIN_DXA   # 9360

BODY_PT = 13.0
LINE_SPACING = 1.15
TITLE_PT = 23.0
SUBTITLE_PT = 11.0
H2_PT = 15.0
SUBLABEL_PT = 10.5
TABLE_HEADER_PT = 9.0
TABLE_BODY_PT = 10.0
CALLOUT_LABEL_PT = 10.5

SECTION_MARKER = "▪ "   # a small filled square, brand blue

# CT_Settings child order (ECMA-376). Used to insert settings elements legally.
SETTINGS_ORDER = [
    "writeProtection", "view", "zoom", "removePersonalInformation",
    "removeDateAndTime", "doNotDisplayPageBoundaries", "displayBackgroundShape",
    "printPostScriptOverText", "printFractionalCharacterWidth", "printFormsData",
    "embedTrueTypeFonts", "embedSystemFonts", "saveSubsetFonts",
    "saveFormsData", "mirrorMargins", "alignBordersAndEdges",
    "bordersDoNotSurroundHeader", "bordersDoNotSurroundFooter", "gutterAtTop",
    "hideSpellingErrors", "hideGrammaticalErrors", "activeWritingStyle",
    "proofState", "formsDesign", "attachedTemplate", "linkStyles",
    "stylePaneFormatFilter", "stylePaneSortMethod", "documentType",
    "mailMerge", "revisionView", "trackChanges", "autoFormatOverride",
    "defaultTabStop", "compat", "rsids", "themeFontLang", "clrSchemeMapping",
]

DASHES = {
    "—": "-",   # em dash
    "–": "-",   # en dash
    "−": "-",   # minus sign
    "‑": "-",   # non-breaking hyphen
    "→": "->",  # right arrow
}


def sanitize(text: Any) -> str:
    """House rule: short hyphens only, never em or en dashes. Arrows as '->'."""
    s = "" if text is None else str(text)
    for bad, good in DASHES.items():
        s = s.replace(bad, good)
    return s


# --------------------------------------------------------------------------
# Low-level OOXML helpers
# --------------------------------------------------------------------------

def _el(tag: str, **attrs) -> Any:
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def set_run_font(run, name: str = FONT) -> None:
    """python-docx only sets w:ascii; set the other three so Word obeys."""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def shade(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(_el("w:shd", val="clear", color="auto", fill=fill_hex))


def cell_borders(cell, bottom: Optional[Tuple[int, str]] = None) -> None:
    """No vertical borders, ever. Optional horizontal hairline at the bottom."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "right", "insideH", "insideV"):
        borders.append(_el("w:" + edge, val="nil"))
    if bottom:
        size, colour = bottom
        borders.append(_el("w:bottom", val="single", sz=size, space=0, color=colour))
    else:
        borders.append(_el("w:bottom", val="nil"))
    tc_pr.append(borders)


def box_borders(cell, size: int, colour: str) -> None:
    """All four sides - used only by the callout box."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        borders.append(_el("w:" + edge, val="single", sz=size, space=0, color=colour))
    tc_pr.append(borders)


def cell_margins(cell, top=90, bottom=90, left=110, right=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        mar.append(_el("w:" + name, w=val, type="dxa"))
    tc_pr.append(mar)


def row_cant_split(row) -> None:
    row._tr.get_or_add_trPr().append(_el("w:cantSplit", val="true"))


def row_is_header(row) -> None:
    row._tr.get_or_add_trPr().append(_el("w:tblHeader", val="true"))


def table_fixed(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        found = tbl_pr.find(qn(tag))
        if found is not None:
            tbl_pr.remove(found)
    tbl_pr.append(_el("w:tblW", w=width_dxa, type="dxa"))
    tbl_pr.append(_el("w:tblLayout", type="fixed"))


def insert_setting(settings_el, new_el) -> None:
    """Insert into w:settings at a schema-legal position."""
    tag = new_el.tag.split("}")[-1]
    if settings_el.find(qn("w:" + tag)) is not None:
        return
    try:
        rank = SETTINGS_ORDER.index(tag)
    except ValueError:
        settings_el.append(new_el)
        return
    for child in settings_el:
        child_tag = child.tag.split("}")[-1]
        child_rank = SETTINGS_ORDER.index(child_tag) if child_tag in SETTINGS_ORDER else 10 ** 6
        if child_rank > rank:
            child.addprevious(new_el)
            return
    settings_el.append(new_el)


# --------------------------------------------------------------------------
# Builder
# --------------------------------------------------------------------------

ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centre": WD_ALIGN_PARAGRAPH.CENTER,
}


class BriefBuilder:
    """Renders a brief model to a styled .docx. No logo, ever."""

    def __init__(self) -> None:
        self.doc = Document()
        self._setup_page()
        self._setup_normal()
        self._setup_view()

    # ---- document chrome -------------------------------------------------

    def _setup_page(self) -> None:
        s = self.doc.sections[0]
        s.page_width = Twips(PAGE_W_DXA)
        s.page_height = Twips(PAGE_H_DXA)
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(s, attr, Twips(MARGIN_DXA))
        s.header_distance = Twips(720)
        s.footer_distance = Twips(720)

    def _setup_normal(self) -> None:
        st = self.doc.styles["Normal"]
        st.font.name = FONT
        st.font.size = Pt(BODY_PT)
        st.font.color.rgb = RGBColor.from_string(GRAPHITE)
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), FONT)
        pf = st.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)   # zero paragraph spacing; blank lines separate blocks

    def _setup_view(self) -> None:
        """Pageless in Google Docs, Web Layout in Word."""
        insert_setting(self.doc.settings.element, _el("w:view", val="web"))

    def enable_font_embedding_flag(self) -> None:
        insert_setting(self.doc.settings.element, _el("w:embedTrueTypeFonts", val="true"))

    # ---- primitives ------------------------------------------------------

    def _para(self, align: str = "left"):
        p = self.doc.add_paragraph()
        p.alignment = ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        pf = p.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        return p

    def _run(self, p, text: str, size: float, colour: str,
             bold: bool = False, italic: bool = False):
        r = p.add_run(sanitize(text))
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(colour)
        r.bold = bold
        r.italic = italic
        set_run_font(r)
        return r

    def blank(self) -> None:
        self._para()

    def pagebreak(self) -> None:
        self._para().add_run().add_break(WD_BREAK.PAGE)

    # ---- blocks ----------------------------------------------------------

    def title(self, text: str, subtitle: Optional[str] = None) -> None:
        p = self._para()
        self._run(p, text, TITLE_PT, NAVY, bold=True)
        if subtitle:
            p2 = self._para()
            self._run(p2, subtitle, SUBTITLE_PT, MUTED)
        self.blank()

    def h2(self, text: str) -> None:
        # No rule or border, above or below. Ever.
        self.blank()
        p = self._para()
        self._run(p, SECTION_MARKER, H2_PT, BRAND_BLUE, bold=True)
        self._run(p, text, H2_PT, NAVY, bold=True)
        self.blank()

    def sublabel(self, text: str) -> None:
        p = self._para()
        self._run(p, text, SUBLABEL_PT, NAVY, bold=True)

    def body(self, text: str) -> None:
        p = self._para()
        self._run(p, text, BODY_PT, GRAPHITE)

    def caption(self, text: str) -> None:
        p = self._para()
        self._run(p, text, SUBLABEL_PT, MUTED)

    def bullets(self, items: Sequence[Any]) -> None:
        """One block. Bold navy lead-in, graphite continuation."""
        for item in items:
            p = self._para()
            p.paragraph_format.left_indent = Twips(360)
            p.paragraph_format.first_line_indent = Twips(-200)
            self._run(p, "•  ", BODY_PT, NAVY, bold=True)
            if isinstance(item, dict):
                lead = item.get("lead", "")
                rest = item.get("text", "")
                if lead:
                    self._run(p, lead, BODY_PT, NAVY, bold=True)
                if rest:
                    sep = "" if (not lead or rest[:1] in " ,.:;") else " "
                    self._run(p, sep + rest, BODY_PT, GRAPHITE)
            else:
                self._run(p, item, BODY_PT, GRAPHITE)

    def fields(self, items: Sequence[Sequence[str]]) -> None:
        """Label / value lines - the spine of a creative brief."""
        for pair in items:
            label = pair[0] if len(pair) > 0 else ""
            value = pair[1] if len(pair) > 1 else ""
            p = self._para()
            self._run(p, f"{label}: ", BODY_PT, NAVY, bold=True)
            self._run(p, value, BODY_PT, GRAPHITE)

    def callout(self, label: str, body: Sequence[str]) -> None:
        t = self.doc.add_table(rows=1, cols=1)
        table_fixed(t, CONTENT_DXA)
        cell = t.cell(0, 0)
        cell.width = Twips(CONTENT_DXA)
        shade(cell, CALLOUT_FILL)
        box_borders(cell, 4, CALLOUT_BORDER)
        cell_margins(cell, top=180, bottom=180, left=220, right=220)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        first = cell.paragraphs[0]
        first.paragraph_format.line_spacing = LINE_SPACING
        first.paragraph_format.space_before = Pt(0)
        first.paragraph_format.space_after = Pt(0)
        if label:
            r = first.add_run(sanitize(label))
            r.font.size = Pt(CALLOUT_LABEL_PT)
            r.font.color.rgb = RGBColor.from_string(BRAND_BLUE)
            r.bold = True
            set_run_font(r)
        target = None if label else first
        for line in body:
            p = target or cell.add_paragraph()
            target = None
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(sanitize(line))
            r.font.size = Pt(BODY_PT)
            r.font.color.rgb = RGBColor.from_string(GRAPHITE)
            set_run_font(r)
        row_cant_split(t.rows[0])
        self.blank()

    def table(self, columns: Sequence[str], rows: Sequence[Sequence[Any]],
              widths: Optional[Sequence[int]] = None,
              align: Optional[Sequence[str]] = None,
              body_pt: float = TABLE_BODY_PT) -> None:
        ncols = len(columns)
        widths = list(widths) if widths else self._even_widths(ncols)
        if sum(widths) != CONTENT_DXA:
            widths = self._rescale(widths)
        align = list(align) if align else ["left"] * ncols

        t = self.doc.add_table(rows=1, cols=ncols)
        table_fixed(t, CONTENT_DXA)

        hdr = t.rows[0]
        row_is_header(hdr)
        row_cant_split(hdr)
        for i, label in enumerate(columns):
            cell = hdr.cells[i]
            cell.width = Twips(widths[i])
            shade(cell, NAVY)
            cell_borders(cell)
            cell_margins(cell, top=90, bottom=80)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            p = cell.paragraphs[0]
            p.alignment = ALIGN.get(align[i], WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(sanitize(label))
            r.font.size = Pt(TABLE_HEADER_PT)
            r.font.color.rgb = RGBColor.from_string(WHITE)
            r.bold = True
            set_run_font(r)

        last = len(rows) - 1
        for ri, data in enumerate(rows):
            tr = t.add_row()
            row_cant_split(tr)
            for ci in range(ncols):
                cell = tr.cells[ci]
                cell.width = Twips(widths[ci])
                cell_borders(cell, bottom=None if ri == last else (4, ROW_HAIRLINE))
                cell_margins(cell, top=110, bottom=110)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                value = data[ci] if ci < len(data) else ""
                lines = sanitize(value).split("\n")
                for li, line in enumerate(lines):
                    p = cell.paragraphs[0] if li == 0 else cell.add_paragraph()
                    p.alignment = ALIGN.get(align[ci], WD_ALIGN_PARAGRAPH.LEFT)
                    p.paragraph_format.line_spacing = LINE_SPACING
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    r = p.add_run(line)
                    r.font.size = Pt(body_pt)
                    r.font.color.rgb = RGBColor.from_string(GRAPHITE)
                    set_run_font(r)
        self.blank()

    @staticmethod
    def _even_widths(n: int) -> List[int]:
        base = CONTENT_DXA // n
        widths = [base] * n
        widths[-1] += CONTENT_DXA - sum(widths)
        return widths

    @staticmethod
    def _rescale(widths: Sequence[int]) -> List[int]:
        total = sum(widths) or 1
        out = [max(300, int(w * CONTENT_DXA / total)) for w in widths]
        out[-1] += CONTENT_DXA - sum(out)
        return out

    # ---- render ----------------------------------------------------------

    def render(self, model: Dict[str, Any]) -> None:
        self.title(model.get("title", "Untitled"), model.get("subtitle"))
        for block in model.get("blocks", []):
            kind = block.get("type", "body")
            if kind == "h2":
                self.h2(block.get("text", ""))
            elif kind == "sublabel":
                self.sublabel(block.get("text", ""))
            elif kind == "body":
                self.body(block.get("text", ""))
            elif kind == "caption":
                self.caption(block.get("text", ""))
            elif kind == "bullets":
                self.bullets(block.get("items", []))
            elif kind == "fields":
                self.fields(block.get("items", []))
            elif kind == "callout":
                self.callout(block.get("label", ""), block.get("body", []))
            elif kind == "table":
                self.table(block.get("columns", []), block.get("rows", []),
                           block.get("widths"), block.get("align"),
                           block.get("body_pt", TABLE_BODY_PT))
            elif kind == "spacer":
                self.blank()
            elif kind == "pagebreak":
                self.pagebreak()
            else:
                raise ValueError(f"unknown block type: {kind!r}")

    def save(self, path: str) -> None:
        self.doc.save(path)


# --------------------------------------------------------------------------
# Markdown dialect -> model
# --------------------------------------------------------------------------

FIELD_RE = re.compile(r"^\*\*(?P<label>[^*]+?):?\*\*:?\s*(?P<value>.*)$")
BULLET_RE = re.compile(r"^[-*]\s+(?P<rest>.*)$")
LEAD_RE = re.compile(r"^\*\*(?P<lead>[^*]+)\*\*(?P<rest>.*)$")
ROW_SEP_RE = re.compile(r"^\|?[\s:|-]+\|[\s:|-]*$")


def _cells(line: str) -> List[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def parse_markdown(text: str) -> Dict[str, Any]:
    lines = text.replace("\r\n", "\n").split("\n")
    model: Dict[str, Any] = {"title": "Untitled", "subtitle": None, "blocks": []}
    blocks = model["blocks"]

    i = 0
    seen_title = False
    pending_fields: List[List[str]] = []
    pending_bullets: List[Any] = []
    pending_body: List[str] = []

    def flush() -> None:
        nonlocal pending_fields, pending_bullets, pending_body
        if pending_fields:
            blocks.append({"type": "fields", "items": pending_fields})
            pending_fields = []
        if pending_bullets:
            blocks.append({"type": "bullets", "items": pending_bullets})
            pending_bullets = []
        if pending_body:
            blocks.append({"type": "body", "text": " ".join(pending_body)})
            pending_body = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            flush()
            i += 1
            continue

        if line.startswith("# ") and not seen_title:
            flush()
            model["title"] = line[2:].strip()
            seen_title = True
            # a single non-blank, non-markup line straight after = subtitle
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if nxt and not nxt.startswith(("#", "-", "*", ">", "|")):
                    model["subtitle"] = nxt.strip("*_ ")
                    i += 1
            i += 1
            continue

        if line.startswith("## "):
            flush()
            blocks.append({"type": "h2", "text": line[3:].strip()})
            i += 1
            continue

        if line.startswith("### "):
            flush()
            blocks.append({"type": "sublabel", "text": line[4:].strip()})
            i += 1
            continue

        if line.startswith(">"):
            flush()
            quoted: List[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quoted.append(lines[i].strip().lstrip(">").strip())
                i += 1
            label = ""
            if quoted and quoted[0].startswith("**") and quoted[0].endswith("**"):
                label = quoted.pop(0).strip("*")
            blocks.append({"type": "callout", "label": label,
                           "body": [q for q in quoted if q]})
            continue

        if line.startswith("|"):
            flush()
            rows: List[List[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                candidate = lines[i].strip()
                if not ROW_SEP_RE.match(candidate):
                    rows.append(_cells(candidate))
                i += 1
            if rows:
                blocks.append({"type": "table", "columns": rows[0], "rows": rows[1:]})
            continue

        if line in ("---", "***", "___"):
            flush()
            blocks.append({"type": "spacer"})
            i += 1
            continue

        m = BULLET_RE.match(line)
        if m:
            if pending_fields or pending_body:
                flush()
            rest = m.group("rest")
            lead_m = LEAD_RE.match(rest)
            if lead_m:
                pending_bullets.append({"lead": lead_m.group("lead"),
                                        "text": lead_m.group("rest")})
            else:
                pending_bullets.append(rest)
            i += 1
            continue

        m = FIELD_RE.match(line)
        if m:
            if pending_bullets or pending_body:
                flush()
            pending_fields.append([m.group("label"), m.group("value")])
            i += 1
            continue

        if pending_fields or pending_bullets:
            flush()
        pending_body.append(line)
        i += 1

    flush()
    return model


def load_model(path: str, fmt: Optional[str] = None) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as fh:
        raw = fh.read()
    kind = fmt or ("json" if path.lower().endswith(".json") else "markdown")
    if kind == "json":
        return json.loads(raw)
    return parse_markdown(raw)


# --------------------------------------------------------------------------
# Font embedding (post-process the package)
# --------------------------------------------------------------------------

FONT_SLOTS = {
    "regular": "embedRegular",
    "bold": "embedBold",
    "semibold": "embedBold",     # Urbanist SemiBold occupies the bold slot
    "italic": "embedItalic",
    "bolditalic": "embedBoldItalic",
    "semibolditalic": "embedBoldItalic",
}


def obfuscate(data: bytes, guid: str) -> bytes:
    """Word's .odttf obfuscation: XOR the first 32 bytes with the reversed GUID."""
    hexes = guid.strip("{}").replace("-", "")
    key = bytes(int(hexes[j:j + 2], 16) for j in range(30, -2, -2))
    head = bytearray(data[:32])
    for idx in range(len(head)):
        head[idx] ^= key[idx % 16]
    return bytes(head) + data[32:]


def discover_fonts(fonts_dir: str) -> Dict[str, str]:
    """Map a style slot -> file path, from filenames like Urbanist-SemiBold.ttf."""
    found: Dict[str, str] = {}
    if not os.path.isdir(fonts_dir):
        return found
    for name in sorted(os.listdir(fonts_dir)):
        if not name.lower().endswith(".ttf"):
            continue
        stem = os.path.splitext(name)[0].lower()
        suffix = stem.split("-")[-1] if "-" in stem else "regular"
        slot = FONT_SLOTS.get(suffix.replace(" ", ""))
        if slot and slot not in found.values():
            found[os.path.join(fonts_dir, name)] = slot
    return found


def embed_fonts(docx_path: str, fonts_dir: str, family: str = FONT) -> int:
    """Embed TTFs into an existing .docx. Returns the number embedded."""
    mapping = discover_fonts(fonts_dir)
    if not mapping:
        return 0

    tmp_dir = tempfile.mkdtemp()
    tmp_out = os.path.join(tmp_dir, "out.docx")
    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            names = zin.namelist()
            parts = {n: zin.read(n) for n in names}

        # 1. [Content_Types].xml - declare the odttf default
        ct = parts["[Content_Types].xml"].decode("utf-8")
        if 'Extension="odttf"' not in ct:
            insert = ('<Default Extension="odttf" ContentType='
                      '"application/vnd.openxmlformats-officedocument.obfuscatedFont"/>')
            head_end = ct.index(">", ct.index("<Types")) + 1
            ct = ct[:head_end] + insert + ct[head_end:]
        parts["[Content_Types].xml"] = ct.encode("utf-8")

        # 2. word/fontTable.xml - register the family with fontKeys
        entries, rels = [], []
        for idx, (path, slot) in enumerate(mapping.items(), start=1):
            rid = f"rIdFont{idx}"
            guid = "{%s}" % str(uuid.uuid4()).upper()
            with open(path, "rb") as fh:
                blob = obfuscate(fh.read(), guid)
            target = f"fonts/font{idx}.odttf"
            parts[f"word/{target}"] = blob
            entries.append(f'<w:{slot} r:id="{rid}" w:fontKey="{guid}" w:subsetted="0"/>')
            rels.append(f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org'
                        f'/officeDocument/2006/relationships/font" Target="{target}"/>')

        font_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<w:font w:name="{family}">'
            '<w:charset w:val="00"/><w:family w:val="swiss"/><w:pitch w:val="variable"/>'
            + "".join(entries) +
            "</w:font></w:fonts>"
        )
        parts["word/fontTable.xml"] = font_xml.encode("utf-8")
        parts["word/_rels/fontTable.xml.rels"] = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            + "".join(rels) + "</Relationships>"
        ).encode("utf-8")

        # 3. document.xml.rels must already reference fontTable; python-docx ships it.
        doc_rels = parts.get("word/_rels/document.xml.rels", b"").decode("utf-8")
        if "fontTable.xml" not in doc_rels:
            doc_rels = doc_rels.replace(
                "</Relationships>",
                '<Relationship Id="rIdFontTable" Type="http://schemas.openxmlformats.org'
                '/officeDocument/2006/relationships/fontTable" Target="fontTable.xml"/>'
                "</Relationships>")
            parts["word/_rels/document.xml.rels"] = doc_rels.encode("utf-8")

        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, blob in parts.items():
                zout.writestr(name, blob)
        shutil.move(tmp_out, docx_path)
        return len(mapping)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# --------------------------------------------------------------------------
# Self-test fixture - exercises every block type
# --------------------------------------------------------------------------

SELF_TEST: Dict[str, Any] = {
    "title": "Sample brand - creative brief",
    "subtitle": "August 2026",
    "blocks": [
        {"type": "callout", "label": "Brief summary",
         "body": ["Three concepts against the problem-aware stage, one desire, one avatar.",
                  "Shoot list and per-line visual direction are in each concept section."]},
        {"type": "h2", "text": "Concept 1 - the reframe"},
        {"type": "sublabel", "text": "Strategy"},
        {"type": "fields", "items": [
            ["Angle", "The routine was never the problem, the timing was"],
            ["Persona", "Shift workers, 28-45, tried three sleep aids"],
            ["Awareness stage", "Problem-aware (canon section 2)"],
            ["Hook", "For anyone who is tired at 3pm and wide awake at 11"],
            ["Format", "Talking head plus b-roll"],
            ["Mechanic", "The reframe"]]},
        {"type": "sublabel", "text": "Script and visual direction"},
        {"type": "table",
         "columns": ["#", "Line", "Visual direction"],
         "widths": [600, 4200, 4560],
         "align": ["center", "left", "left"],
         "rows": [
             ["1", "For anyone who is tired at 3pm and wide awake at 11.",
              "Talking head, kitchen, morning light"],
             ["2", "It was never your routine. It was when you started it.",
              "Close-up of the clock, then the product on the counter"],
             ["3", "Click below to get your first week for the price of a coffee.",
              "Product in hand, walking out of frame"]]},
        {"type": "sublabel", "text": "Production notes"},
        {"type": "bullets", "items": [
            {"lead": "Talent", "text": "cast for resemblance, not for fitness."},
            {"lead": "Environment", "text": "real kitchen, no studio."},
            "Deliver 9:16 and 4:5. Captions burned in, varied per scene."]},
        {"type": "spacer"},
        {"type": "body",
         "text": "Metric accountability: hook rate and hold rate for the fast read; "
                 "CPA with model and window stated for the scale decision."},
        {"type": "caption", "text": "Deltas colour-coded per the document style spec."},
    ],
}


def main(argv: Optional[Sequence[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="Build a VIDEN-styled brief .docx")
    ap.add_argument("--in", dest="src", help="brief.json or brief.md")
    ap.add_argument("--out", dest="out", required=True, help="output .docx path")
    ap.add_argument("--format", choices=["json", "markdown"], default=None)
    ap.add_argument("--fonts", dest="fonts", help="directory of Urbanist TTFs")
    ap.add_argument("--embed", action="store_true", help="embed the TTFs in --fonts")
    ap.add_argument("--self-test", action="store_true",
                    help="build the built-in fixture instead of reading --in")
    args = ap.parse_args(argv)

    if args.self_test:
        model = SELF_TEST
    elif args.src:
        model = load_model(args.src, args.format)
    else:
        ap.error("give --in PATH or --self-test")
        return 2

    builder = BriefBuilder()
    if args.embed:
        builder.enable_font_embedding_flag()
    builder.render(model)
    builder.save(args.out)

    embedded = 0
    if args.embed:
        if not args.fonts:
            print("warning: --embed given without --fonts; skipping embedding",
                  file=sys.stderr)
        else:
            embedded = embed_fonts(args.out, args.fonts)
            if embedded == 0:
                print(f"warning: no .ttf files found in {args.fonts}; skipping embedding",
                      file=sys.stderr)

    blocks = len(model.get("blocks", []))
    size = os.path.getsize(args.out)
    print(f"wrote {args.out} ({size} bytes, {blocks} blocks"
          f"{f', {embedded} fonts embedded' if embedded else ''})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
